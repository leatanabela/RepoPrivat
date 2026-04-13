import os
import pdfplumber
from docx import Document


def _ocr_pdf(file_path: str) -> str:
    """Fallback: use OCR to extract text from scanned PDF pages."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        raise ValueError(
            "PDF-ul pare scanat (fara text selectabil). "
            "Instaleaza OCR: pip install pdf2image pytesseract && brew install tesseract poppler"
        )

    images = convert_from_path(file_path, dpi=300)
    text_parts = []
    for img in images:
        page_text = pytesseract.image_to_string(img, lang="ron+eng")
        if page_text.strip():
            text_parts.append(page_text.strip())
    return "\n\n".join(text_parts)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file. Falls back to OCR for scanned PDFs."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    text = "\n\n".join(text_parts)

    # If pdfplumber got almost no text, try OCR
    if len(text.strip()) < 100:
        return _ocr_pdf(file_path)

    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = Document(file_path)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


def extract_text(file_path: str) -> str:
    """Auto-detect file type and extract text."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")
